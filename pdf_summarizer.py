import gradio as gr
from langchain_community.chat_models import ChatOpenAI
from langchain.schema import HumanMessage, SystemMessage
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains.summarize import load_summarize_chain
from langchain.docstore.document import Document
from langchain.prompts import PromptTemplate
import PyPDF2
import docx
import io
import warnings
import re
from typing import List, Tuple, Dict, Optional
import nltk
from nltk.tokenize import sent_tokenize
import os
import tempfile
import pdfplumber
import platform
import logging
import hashlib
import json
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed, TimeoutError
import time
from functools import lru_cache
import threading
from queue import Queue
import gc
import psutil
import asyncio
from datetime import datetime
import requests

# Suppress warnings
logging.getLogger("langchain.text_splitter").setLevel(logging.ERROR)
logging.getLogger("langchain_community.chat_models.openai").setLevel(logging.ERROR)

# Try to import OCR dependencies
OCR_AVAILABLE = False
try:
    import pytesseract
    from pdf2image import convert_from_path, convert_from_bytes
    from PIL import Image
    import numpy as np
    import cv2

    # Configure Tesseract path for Windows
    if platform.system() == 'Windows':
        tesseract_paths = [
            r"D:\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            r"C:\Tesseract-OCR\tesseract.exe",
        ]

        for path in tesseract_paths:
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                break

    OCR_AVAILABLE = True
except ImportError:
    print("⚠️ OCR dependencies not installed. OCR features will be disabled.")
    print("⚠️ 未安装OCR依赖项。OCR功能将被禁用。")

warnings.filterwarnings('ignore')

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    print("Downloading NLTK punkt tokenizer...")
    nltk.download('punkt', quiet=True)


class ModelConfig:
    """Configuration for different AI models"""

    MODELS = {
        "deepseek-chat": {
            "name": "🚀 DeepSeek Chat",
            "provider": "deepseek",
            "api_base": "https://api.deepseek.com",
            "model_name": "deepseek-chat",
            "max_tokens": 1024,
            "temperature": 0.3,
            "supports_streaming": True,
            "api_key_placeholder": "输入您的DeepSeek API密钥... Enter your DeepSeek API key...",
            "api_key_help": "获取API密钥：https://platform.deepseek.com/api_keys"
        },
        "doubao-pro-4k": {
            "name": "🌊 豆包 Pro 4K (ByteDance)",
            "provider": "doubao",
            "api_base": "https://ark.cn-beijing.volces.com/api/v3",
            "model_name": "doubao-pro-4k",
            "max_tokens": 1024,
            "temperature": 0.3,
            "supports_streaming": True,
            "api_key_placeholder": "输入您的豆包API密钥... Enter your Doubao API key...",
            "api_key_help": "获取API密钥：https://console.volcengine.com/ark"
        },
        "doubao-lite-4k": {
            "name": "🌊 豆包 Lite 4K (ByteDance)",
            "provider": "doubao",
            "api_base": "https://ark.cn-beijing.volces.com/api/v3",
            "model_name": "doubao-lite-4k",
            "max_tokens": 1024,
            "temperature": 0.3,
            "supports_streaming": True,
            "api_key_placeholder": "输入您的豆包API密钥... Enter your Doubao API key...",
            "api_key_help": "获取API密钥：https://console.volcengine.com/ark"
        },
        "ollama-llama3": {
            "name": "🦙 Ollama Llama 3",
            "provider": "ollama",
            "api_base": "http://localhost:11434",
            "model_name": "llama3",
            "max_tokens": 1024,
            "temperature": 0.3,
            "supports_streaming": False,
            "api_key_placeholder": "Ollama本地运行，无需API密钥 Ollama runs locally, no API key needed",
            "api_key_help": "确保Ollama正在运行：ollama serve"
        },
        "ollama-qwen": {
            "name": "🦙 Ollama Qwen",
            "provider": "ollama",
            "api_base": "http://localhost:11434",
            "model_name": "qwen",
            "max_tokens": 1024,
            "temperature": 0.3,
            "supports_streaming": False,
            "api_key_placeholder": "Ollama本地运行，无需API密钥 Ollama runs locally, no API key needed",
            "api_key_help": "确保Ollama正在运行：ollama serve"
        },
        "ollama-mistral": {
            "name": "🦙 Ollama Mistral",
            "provider": "ollama",
            "api_base": "http://localhost:11434",
            "model_name": "mistral",
            "max_tokens": 1024,
            "temperature": 0.3,
            "supports_streaming": False,
            "api_key_placeholder": "Ollama本地运行，无需API密钥 Ollama runs locally, no API key needed",
            "api_key_help": "确保Ollama正在运行：ollama serve"
        }
    }

    @classmethod
    def get_model_choices(cls):
        """Get model choices for UI dropdown"""
        return [(config["name"], model_id) for model_id, config in cls.MODELS.items()]

    @classmethod
    def get_config(cls, model_id):
        """Get configuration for a specific model"""
        return cls.MODELS.get(model_id)


class OllamaClient:
    """Custom Ollama client for local models"""

    def __init__(self, base_url="http://localhost:11434", model_name="llama3"):
        self.base_url = base_url.rstrip('/')
        self.model_name = model_name

    def _make_request(self, endpoint, data):
        """Make HTTP request to Ollama API"""
        url = f"{self.base_url}/{endpoint}"
        try:
            response = requests.post(url, json=data, timeout=120, stream=False)
            response.raise_for_status()
            return response.json()
        except requests.exceptions.RequestException as e:
            raise Exception(f"Ollama API error: {str(e)}")

    def generate(self, prompt, max_tokens=1024, temperature=0.3):
        """Generate text using Ollama"""
        data = {
            "model": self.model_name,
            "prompt": prompt,
            "options": {
                "num_predict": max_tokens,
                "temperature": temperature
            },
            "stream": False
        }

        try:
            result = self._make_request("api/generate", data)
            return result.get("response", "")
        except Exception as e:
            return f"Ollama生成错误 Ollama generation error: {str(e)}"

    def chat(self, messages, max_tokens=1024, temperature=0.3):
        """Chat with Ollama using message format"""
        # Convert messages to a single prompt
        prompt_parts = []
        for msg in messages:
            if hasattr(msg, 'content'):
                content = msg.content
            else:
                content = str(msg)

            if hasattr(msg, 'type'):
                if msg.type == 'system':
                    prompt_parts.append(f"System: {content}")
                elif msg.type == 'human':
                    prompt_parts.append(f"Human: {content}")
                else:
                    prompt_parts.append(content)
            else:
                prompt_parts.append(content)

        prompt = "\n\n".join(prompt_parts) + "\n\nAssistant:"
        return self.generate(prompt, max_tokens, temperature)


class DocumentCache:
    """Simple cache for processed documents"""

    def __init__(self, cache_dir=None):
        self.cache_dir = cache_dir or tempfile.gettempdir()
        self.cache_path = Path(self.cache_dir) / "doc_summarizer_cache"
        self.cache_path.mkdir(exist_ok=True)
        self.cache_index = self.cache_path / "index.json"
        self.load_index()

    def load_index(self):
        """Load cache index"""
        if self.cache_index.exists():
            try:
                with open(self.cache_index, 'r') as f:
                    self.index = json.load(f)
            except:
                self.index = {}
        else:
            self.index = {}

    def save_index(self):
        """Save cache index"""
        try:
            with open(self.cache_index, 'w') as f:
                json.dump(self.index, f)
        except:
            pass

    def get_file_hash(self, file_path):
        """Get hash of file for cache key"""
        try:
            hasher = hashlib.md5()
            with open(file_path, 'rb') as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hasher.update(chunk)
            return hasher.hexdigest()
        except:
            return None

    def get(self, file_path, cache_type="text"):
        """Get cached result if exists"""
        try:
            file_hash = self.get_file_hash(file_path)
            if not file_hash:
                return None

            cache_key = f"{file_hash}_{cache_type}"

            if cache_key in self.index:
                cache_file = self.cache_path / self.index[cache_key]
                if cache_file.exists():
                    with open(cache_file, 'r', encoding='utf-8') as f:
                        return f.read()
        except:
            pass
        return None

    def set(self, file_path, content, cache_type="text"):
        """Cache result"""
        try:
            file_hash = self.get_file_hash(file_path)
            if not file_hash:
                return

            cache_key = f"{file_hash}_{cache_type}"
            cache_file = self.cache_path / f"{cache_key}.txt"

            with open(cache_file, 'w', encoding='utf-8') as f:
                f.write(content)

            self.index[cache_key] = cache_file.name
            self.save_index()
        except:
            pass


class OptimizedDocumentSummarizer:
    def __init__(self, model_id="deepseek-chat", api_key=None):
        self.model_id = model_id
        self.model_config = ModelConfig.get_config(model_id)
        self.api_key = api_key

        if not self.model_config:
            raise ValueError(f"Unsupported model: {model_id}")

        # Initialize the appropriate client
        self._initialize_client()

        # Text splitters with smaller chunks
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", "。", ". ", "！", "! ", "？", "? ", "；", "; ", " ", ""],
            is_separator_regex=False
        )

        # Initialize cache
        self.cache = DocumentCache()

        # Thread pool for parallel processing
        self.executor = ThreadPoolExecutor(max_workers=2)

        # OCR configuration
        self.ocr_available = False
        self.chinese_ocr_available = False
        if OCR_AVAILABLE:
            self.configure_ocr()

        # Add cancellation token
        self.cancel_processing = False

        # Maximum text length to process (characters)
        self.max_text_length = 200000

        # Maximum chunks to process
        self.max_chunks = 20

    def _initialize_client(self):
        """Initialize the appropriate client based on model provider"""
        provider = self.model_config["provider"]

        if provider == "ollama":
            # Initialize Ollama client
            self.llm = None  # Will use custom client
            self.ollama_client = OllamaClient(
                base_url=self.model_config["api_base"],
                model_name=self.model_config["model_name"]
            )
        else:
            # Initialize OpenAI-compatible client (DeepSeek, Doubao)
            self.llm = ChatOpenAI(
                model=self.model_config["model_name"],
                openai_api_key=self.api_key or "dummy-key",
                openai_api_base=self.model_config["api_base"],
                max_tokens=self.model_config["max_tokens"],
                temperature=self.model_config["temperature"],
                streaming=self.model_config["supports_streaming"],
                request_timeout=60
            )
            self.ollama_client = None

    def test_connection(self):
        """Test connection to the model"""
        try:
            if self.model_config["provider"] == "ollama":
                # Test Ollama connection
                test_response = self.ollama_client.generate("Test", max_tokens=10)
                if "error" in test_response.lower():
                    return False, test_response
                return True, "Ollama连接成功 Ollama connection successful"
            else:
                # Test OpenAI-compatible API
                test_messages = [
                    SystemMessage(content="You are a helpful assistant."),
                    HumanMessage(content="Hello")
                ]
                response = self.llm.invoke(test_messages)
                return True, f"{self.model_config['name']} 连接成功 connection successful"
        except Exception as e:
            return False, f"连接失败 Connection failed: {str(e)}"

    def configure_ocr(self):
        """Configure OCR settings"""
        try:
            version = pytesseract.get_tesseract_version()
            self.ocr_available = True
            available_langs = pytesseract.get_languages()
            self.chinese_ocr_available = any(lang in available_langs for lang in ['chi_sim', 'chi_tra'])
        except:
            self.ocr_available = False
            self.chinese_ocr_available = False

    def extract_text_from_pdf_fast(self, file_path, use_ocr_if_needed=True, ocr_language='auto',
                                   quality='balanced', progress_callback=None, max_ocr_pages=20):
        """Optimized PDF extraction with timeout and page limits"""

        # Reset cancellation token
        self.cancel_processing = False

        # Check cache first
        cache_key = f"{quality}_{ocr_language}_ocr{use_ocr_if_needed}_max{max_ocr_pages}"
        cached_text = self.cache.get(file_path, cache_key)
        if cached_text:
            if progress_callback:
                progress_callback(1.0, "从缓存加载 Loaded from cache")
            return cached_text

        extracted_text = ""
        ocr_pages = []

        # Try fast extraction with pdfplumber first
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)

                # Limit pages for very large PDFs
                max_pages_to_extract = min(total_pages, 100)

                if total_pages > max_pages_to_extract:
                    if progress_callback:
                        progress_callback(0.05,
                                          f"大型PDF检测到：仅处理前{max_pages_to_extract}页 Large PDF detected: Processing first {max_pages_to_extract} pages only")

                # Quick scan to check if OCR is needed (only check first 3 pages)
                sample_pages = min(3, total_pages)
                needs_ocr = False

                for i in range(sample_pages):
                    if self.cancel_processing:
                        return "用户已取消处理。Processing cancelled by user."

                    page_text = pdf.pages[i].extract_text() or ""
                    if self.is_scanned_pdf_page(page_text) or self.is_text_corrupted(page_text):
                        needs_ocr = True
                        break

                if not needs_ocr:
                    # Extract text with length limit
                    for i in range(max_pages_to_extract):
                        if self.cancel_processing:
                            return "用户已取消处理。Processing cancelled by user."

                        if len(extracted_text) > self.max_text_length:
                            extracted_text += f"\n\n--- 达到文本长度限制，停止提取 Text length limit reached, stopping extraction ---\n"
                            break

                        if progress_callback:
                            progress_callback(i / max_pages_to_extract,
                                              f"提取第 {i + 1}/{max_pages_to_extract} 页 Extracting page {i + 1}/{max_pages_to_extract}")

                        try:
                            page_text = pdf.pages[i].extract_text() or ""
                            if page_text and not self.is_text_corrupted(page_text):
                                extracted_text += f"\n--- 第 {i + 1}/{total_pages} 页 Page {i + 1}/{total_pages} ---\n{page_text}\n"
                        except Exception as e:
                            print(f"Error extracting page {i + 1}: {str(e)}")
                            continue

                    if extracted_text:
                        # Truncate if still too long
                        if len(extracted_text) > self.max_text_length:
                            extracted_text = extracted_text[:self.max_text_length] + "\n\n--- 文本已截断 Text truncated ---"

                        self.cache.set(file_path, extracted_text, cache_key)
                        return extracted_text
        except Exception as e:
            print(f"PDFPlumber error: {str(e)}")

        # If fast extraction failed or OCR is needed
        if use_ocr_if_needed and self.ocr_available:
            return self._extract_with_limited_ocr(file_path, ocr_language, quality, progress_callback, max_ocr_pages)
        else:
            # Fallback to PyPDF2
            return self._extract_with_pypdf2(file_path, progress_callback)

    def _extract_with_limited_ocr(self, file_path, ocr_language, quality, progress_callback, max_ocr_pages):
        """Extract text using OCR with page limits and timeout"""

        # Determine DPI based on quality setting
        dpi_settings = {
            'fast': 100,
            'balanced': 150,
            'high': 200
        }
        dpi = dpi_settings.get(quality, 150)

        try:
            # First, get total page count
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)

            if progress_callback:
                progress_callback(0.1,
                                  f"PDF共有 {total_pages} 页，将OCR处理前 {max_ocr_pages} 页... PDF has {total_pages} pages. Will OCR up to {max_ocr_pages} pages...")

            # Limit pages to process
            pages_to_process = min(total_pages, max_ocr_pages)

            # Convert only the pages we need
            extracted_text = ""

            # Process pages in smaller batches
            batch_size = 5

            for batch_start in range(0, pages_to_process, batch_size):
                if self.cancel_processing:
                    return "用户已取消处理。Processing cancelled by user."

                if len(extracted_text) > self.max_text_length:
                    extracted_text += f"\n\n--- 达到文本长度限制 Text length limit reached ---\n"
                    break

                batch_end = min(batch_start + batch_size, pages_to_process)

                # Convert batch of pages
                try:
                    if progress_callback:
                        progress_callback(
                            0.1 + (0.8 * batch_start / pages_to_process),
                            f"转换第 {batch_start + 1}-{batch_end} 页... Converting pages {batch_start + 1}-{batch_end}..."
                        )

                    # Convert specific page range
                    images = convert_from_path(
                        file_path,
                        dpi=dpi,
                        first_page=batch_start + 1,
                        last_page=batch_end,
                        thread_count=2,
                        fmt='jpeg',
                        jpegopt={'quality': 75, 'optimize': True}
                    )

                    # Process batch
                    for i, image in enumerate(images):
                        if self.cancel_processing:
                            return "用户已取消处理。Processing cancelled by user."

                        page_num = batch_start + i + 1

                        if progress_callback:
                            progress_callback(
                                0.1 + (0.8 * page_num / pages_to_process),
                                f"OCR处理第 {page_num}/{pages_to_process} 页... OCR processing page {page_num}/{pages_to_process}..."
                            )

                        try:
                            # Process with timeout
                            text = self._ocr_with_timeout(image, ocr_language, quality, timeout=30)
                            if text and text != "OCR timeout" and text != "OCR Error":
                                extracted_text += f"\n--- 第 {page_num}/{total_pages} 页 Page {page_num}/{total_pages} (OCR) ---\n{text}\n"
                        except Exception as e:
                            print(f"处理第 {page_num} 页时出错 Error processing page {page_num}: {str(e)}")
                            extracted_text += f"\n--- 第 {page_num}/{total_pages} 页 Page {page_num}/{total_pages} (OCR失败 Failed) ---\n[此页OCR失败 OCR failed for this page]\n"

                    # Clear memory after each batch
                    del images
                    gc.collect()

                except Exception as e:
                    print(
                        f"转换批次 {batch_start}-{batch_end} 时出错 Error converting batch {batch_start}-{batch_end}: {str(e)}")
                    continue

            # Add note about remaining pages if any
            if total_pages > pages_to_process:
                extracted_text += f"\n\n--- 注意：OCR仅处理了前 {pages_to_process} 页，共 {total_pages} 页 Note: OCR processed first {pages_to_process} pages out of {total_pages} total pages ---\n"

            # Truncate if too long
            if len(extracted_text) > self.max_text_length:
                extracted_text = extracted_text[:self.max_text_length] + "\n\n--- 文本已截断 Text truncated ---"

            # Cache the result
            cache_key = f"{quality}_{ocr_language}_ocrTrue_max{max_ocr_pages}"
            self.cache.set(file_path, extracted_text, cache_key)

            return extracted_text if extracted_text else "无法通过OCR提取文本。No text could be extracted with OCR."

        except Exception as e:
            return f"OCR处理时出错 Error during OCR processing: {str(e)}"

    def _ocr_with_timeout(self, image, ocr_language, quality, timeout=30):
        """Run OCR with timeout"""
        import signal
        from contextlib import contextmanager

        try:
            # Use threading for timeout
            result = [None]
            exception = [None]

            def run_ocr():
                try:
                    result[0] = self.extract_text_with_ocr(
                        image,
                        preprocess=(quality == 'high'),
                        language=ocr_language
                    )
                except Exception as e:
                    exception[0] = e

            thread = threading.Thread(target=run_ocr)
            thread.daemon = True
            thread.start()
            thread.join(timeout)

            if thread.is_alive():
                # OCR is still running after timeout
                return "OCR超时 OCR timeout"

            if exception[0]:
                raise exception[0]

            return result[0] or "未提取到文本 No text extracted"

        except Exception as e:
            return f"OCR错误 OCR Error: {str(e)}"

    def _extract_with_pypdf2(self, file_path, progress_callback):
        """Fallback extraction using PyPDF2"""
        try:
            extracted_text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)

                # Limit pages
                max_pages = min(total_pages, 100)

                for i in range(max_pages):
                    if self.cancel_processing:
                        return "用户已取消处理。Processing cancelled by user."

                    if len(extracted_text) > self.max_text_length:
                        extracted_text += f"\n\n--- 达到文本长度限制 Text length limit reached ---\n"
                        break

                    if progress_callback:
                        progress_callback(i / max_pages,
                                          f"提取第 {i + 1}/{max_pages} 页 Extracting page {i + 1}/{max_pages}")

                    try:
                        page_text = pdf_reader.pages[i].extract_text()
                        if page_text and not self.is_text_corrupted(page_text):
                            extracted_text += f"\n--- 第 {i + 1}/{total_pages} 页 Page {i + 1}/{total_pages} ---\n{page_text}\n"
                    except:
                        continue

            # Truncate if too long
            if len(extracted_text) > self.max_text_length:
                extracted_text = extracted_text[:self.max_text_length] + "\n\n--- 文本已截断 Text truncated ---"

            return extracted_text if extracted_text else "无法从PDF中提取文本。No text could be extracted from the PDF."
        except Exception as e:
            return f"读取PDF时出错 Error reading PDF: {str(e)}"

    def preprocess_image_for_ocr(self, image):
        """Optimized image preprocessing"""
        if not OCR_AVAILABLE:
            return image

        try:
            # Resize if too large
            max_dimension = 2000
            if image.width > max_dimension or image.height > max_dimension:
                ratio = max_dimension / max(image.width, image.height)
                new_size = (int(image.width * ratio), int(image.height * ratio))
                image = image.resize(new_size, Image.Resampling.LANCZOS)

            # Convert to numpy array
            img_array = np.array(image)

            # Quick grayscale conversion
            if len(img_array.shape) == 3:
                gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            else:
                gray = img_array

            # Simple thresholding
            _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

            return Image.fromarray(thresh)
        except:
            return image

    def extract_text_with_ocr(self, image, preprocess=True, language='auto'):
        """Optimized OCR extraction"""
        if not self.ocr_available or not OCR_AVAILABLE:
            return "OCR不可用。OCR not available."

        try:
            if preprocess:
                image = self.preprocess_image_for_ocr(image)

            # Determine OCR language
            ocr_lang = 'eng'
            if language == 'chinese' and self.chinese_ocr_available:
                ocr_lang = 'chi_sim+eng'
            elif language == 'auto':
                ocr_lang = 'eng'

            # Perform OCR with optimized settings
            text = pytesseract.image_to_string(
                image,
                lang=ocr_lang,
                config='--psm 3 --oem 1 -c tessedit_do_invert=0'
            )

            return text.strip()
        except Exception as e:
            return f"OCR错误 OCR Error: {str(e)}"

    def is_scanned_pdf_page(self, page_text):
        """Quick check if page is scanned"""
        return len(page_text.strip()) < 50

    def is_text_corrupted(self, text):
        """Quick corruption check"""
        if not text or len(text.strip()) < 10:
            return True

        # Quick check for readable characters
        readable_chars = len(re.findall(r'[\u4e00-\u9fff\u0020-\u007E\u00A0-\u00FF]', text[:100]))
        return readable_chars < 30

    def extract_text_from_docx(self, file_path):
        """Optimized Word document extraction"""
        # Check cache
        cached_text = self.cache.get(file_path, "docx")
        if cached_text:
            return cached_text

        try:
            doc = docx.Document(file_path)
            text_parts = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if len("".join(text_parts)) > self.max_text_length:
                    text_parts.append("\n\n--- 达到文本长度限制 Text length limit reached ---")
                    break

                if paragraph.text.strip():
                    if paragraph.style and paragraph.style.name.startswith('Heading'):
                        text_parts.append(f"\n## {paragraph.text}\n")
                    else:
                        text_parts.append(paragraph.text)

            # Extract tables efficiently
            for table in doc.tables:
                if len("".join(text_parts)) > self.max_text_length:
                    break

                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_parts.append(row_text)

            text = "\n".join(text_parts)

            # Truncate if too long
            if len(text) > self.max_text_length:
                text = text[:self.max_text_length] + "\n\n--- 文本已截断 Text truncated ---"

            # Cache result
            self.cache.set(file_path, text, "docx")

            return text.strip() if text else "文档中未找到文本。No text found in the document."
        except Exception as e:
            return f"读取Word文档时出错 Error reading Word document: {str(e)}"

    def get_file_text(self, file_path, ocr_language='auto', quality='balanced',
                      progress_callback=None, max_ocr_pages=20):
        """Extract text with progress tracking"""
        file_lower = file_path.lower()

        if file_lower.endswith('.pdf'):
            return self.extract_text_from_pdf_fast(
                file_path,
                ocr_language=ocr_language,
                quality=quality,
                progress_callback=progress_callback,
                max_ocr_pages=max_ocr_pages
            )
        elif file_lower.endswith(('.docx', '.doc')):
            return self.extract_text_from_docx(file_path)
        elif file_lower.endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif')):
            # Check cache
            cached_text = self.cache.get(file_path, f"image_{ocr_language}")
            if cached_text:
                return cached_text

            # Extract from image
            try:
                image = Image.open(file_path)
                text = self.extract_text_with_ocr(image, language=ocr_language)

                # Cache result
                self.cache.set(file_path, text, f"image_{ocr_language}")

                return text if text else "无法从图片中提取文本。No text could be extracted from the image."
            except Exception as e:
                return f"读取图片时出错 Error reading image: {str(e)}"
        elif file_lower.endswith('.txt'):
            try:
                encodings = ['utf-8', 'gbk', 'gb2312', 'big5', 'utf-16']
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as file:
                            text = file.read()
                            # Limit text length
                            if len(text) > self.max_text_length:
                                text = text[:self.max_text_length] + "\n\n--- 文本已截断 Text truncated ---"
                            return text
                    except UnicodeDecodeError:
                        continue
                return "错误：无法解码文本文件。Error: Unable to decode text file."
            except Exception as e:
                return f"读取文本文件时出错 Error reading text file: {str(e)}"
        else:
            return "不支持的文件格式。Unsupported file format."

    def summarize_text_streaming(self, text, summary_type="concise", include_quotes=False,
                                 output_language="auto", progress_callback=None):
        """Generate summary with streaming support and timeout"""

        if not text or text.startswith("Error") or text.startswith("❌"):
            return text

        # Check cache for summary
        cache_key = f"summary_{summary_type}_{include_quotes}_{output_language}_{self.model_id}"
        text_hash = hashlib.md5(text.encode()).hexdigest()
        cached_summary = self.cache.get(text_hash, cache_key)
        if cached_summary:
            return cached_summary

        # Limit text length for summarization
        if len(text) > self.max_text_length:
            text = text[:self.max_text_length]
            if progress_callback:
                progress_callback(0.2, "文本过长，已截断 Text too long, truncated")

        # Create documents with limited chunks
        chunks = self.text_splitter.split_text(text)

        # Limit number of chunks
        if len(chunks) > self.max_chunks:
            chunks = chunks[:self.max_chunks]
            if progress_callback:
                progress_callback(0.3,
                                  f"文档块过多，仅处理前{self.max_chunks}块 Too many chunks, processing first {self.max_chunks}")

        documents = [Document(page_content=chunk) for chunk in chunks]

        if not documents:
            return "未找到可总结的内容。No content found to summarize."

        # Generate summary with timeout protection
        try:
            summary = self._generate_summary_with_timeout(
                documents, summary_type, include_quotes,
                output_language, progress_callback,
                timeout=300  # 5 minute timeout for entire summarization
            )

            # Cache result
            if summary and not summary.startswith("Error") and not summary.startswith("超时"):
                self.cache.set(text_hash, summary, cache_key)

            return summary
        except Exception as e:
            return f"总结生成失败 Summarization failed: {str(e)}"

    def _generate_summary_with_timeout(self, documents, summary_type, include_quotes,
                                       output_language, progress_callback, timeout=300):
        """Generate summary with timeout protection"""

        result = [None]
        exception = [None]

        def run_summary():
            try:
                result[0] = self._generate_summary(
                    documents, summary_type, include_quotes,
                    output_language, progress_callback
                )
            except Exception as e:
                exception[0] = e

        thread = threading.Thread(target=run_summary)
        thread.daemon = True
        thread.start()
        thread.join(timeout)

        if thread.is_alive():
            self.cancel_processing = True
            return "超时：总结生成时间过长，请尝试减少文档大小或使用'简洁'模式 Timeout: Summary generation took too long. Try reducing document size or using 'concise' mode."

        if exception[0]:
            raise exception[0]

        return result[0] or "未能生成摘要 Failed to generate summary"

    def _generate_summary(self, documents, summary_type, include_quotes, output_language, progress_callback):
        """Generate summary with appropriate method"""

        # Language instructions
        lang_instructions = {
            "chinese": "用中文撰写摘要。Write the summary in Chinese (中文).",
            "english": "用英文撰写摘要。Write the summary in English.",
            "auto": "使用源文档的语言撰写摘要。Match the language of the source document."
        }
        lang_instruction = lang_instructions.get(output_language, lang_instructions["auto"])

        # Summary prompts (simplified for better performance)
        prompts = {
            "concise": f"简洁总结以下内容（2-3段）。Concisely summarize (2-3 paragraphs). {lang_instruction}\n\n{{text}}\n\n摘要 SUMMARY:",
            "detailed": f"详细总结以下内容。Provide detailed summary. {lang_instruction}\n\n{{text}}\n\n详细摘要 DETAILED SUMMARY:",
            "bullet_points": f"用要点总结。Summarize in bullet points. {lang_instruction}\n\n{{text}}\n\n要点 BULLET POINTS:",
            "key_insights": f"提取5个关键见解。Extract 5 key insights. {lang_instruction}\n\n{{text}}\n\n关键见解 KEY INSIGHTS:",
            "chapter_wise": f"按章节总结。Summarize by sections. {lang_instruction}\n\n{{text}}\n\n章节摘要 SECTION SUMMARY:"
        }

        prompt_template = prompts.get(summary_type, prompts["concise"])

        try:
            if len(documents) > 1:
                # Use simpler approach for multiple documents
                if progress_callback:
                    progress_callback(0.3,
                                      f"处理 {len(documents)} 个文档块... Processing {len(documents)} document chunks...")

                # Combine all documents first (faster than map-reduce for moderate sizes)
                if len(documents) <= 5:
                    combined_text = "\n\n".join([doc.page_content for doc in documents])

                    # Generate summary based on model type
                    if self.model_config["provider"] == "ollama":
                        # Use Ollama client
                        summary = self._generate_with_ollama(prompt_template.format(text=combined_text),
                                                             progress_callback)
                    else:
                        # Use OpenAI-compatible client with streaming
                        messages = [
                            SystemMessage(content="你是专业的文档分析师。You are a professional document analyst."),
                            HumanMessage(content=prompt_template.format(text=combined_text))
                        ]

                        summary = ""
                        chunk_count = 0

                        try:
                            if self.model_config["supports_streaming"]:
                                for chunk in self.llm.stream(messages):
                                    if self.cancel_processing:
                                        return "用户已取消处理。Processing cancelled by user."

                                    summary += chunk.content
                                    chunk_count += 1

                                    if progress_callback and chunk_count % 10 == 0:
                                        progress_callback(0.5 + 0.4 * min(chunk_count / 100, 1),
                                                          "生成摘要中... Generating summary...")
                            else:
                                response = self.llm.invoke(messages)
                                summary = response.content
                        except Exception as e:
                            return f"API调用失败 API call failed: {str(e)}"

                else:
                    # For larger documents, process in batches
                    batch_summaries = []
                    batch_size = 3

                    for i in range(0, len(documents), batch_size):
                        if self.cancel_processing:
                            return "用户已取消处理。Processing cancelled by user."

                        batch = documents[i:i + batch_size]
                        batch_text = "\n\n".join([doc.page_content for doc in batch])

                        if progress_callback:
                            progress_callback(0.3 + 0.4 * (i / len(documents)),
                                              f"处理批次 {i // batch_size + 1}/{(len(documents) + batch_size - 1) // batch_size}... Processing batch {i // batch_size + 1}/{(len(documents) + batch_size - 1) // batch_size}...")

                        batch_prompt = f"{lang_instruction}\n\n{batch_text}\n\n摘要 SUMMARY:"

                        try:
                            if self.model_config["provider"] == "ollama":
                                batch_summary = self._generate_with_ollama(batch_prompt)
                            else:
                                messages = [
                                    SystemMessage(content="总结这部分内容。Summarize this section."),
                                    HumanMessage(content=batch_prompt)
                                ]
                                batch_summary = self.llm.invoke(messages).content

                            batch_summaries.append(batch_summary)
                        except Exception as e:
                            print(f"批次处理失败 Batch processing failed: {str(e)}")
                            continue

                    # Combine batch summaries
                    if batch_summaries:
                        combined_summaries = "\n\n".join(batch_summaries)
                        final_prompt = prompt_template.format(text=combined_summaries)

                        if self.model_config["provider"] == "ollama":
                            summary = self._generate_with_ollama(final_prompt)
                        else:
                            final_messages = [
                                SystemMessage(
                                    content="合并以下摘要为最终摘要。Combine these summaries into a final summary."),
                                HumanMessage(content=final_prompt)
                            ]
                            summary = self.llm.invoke(final_messages).content
                    else:
                        return "无法生成摘要 Failed to generate summary"

            else:
                # Direct summarization for single document
                if self.model_config["provider"] == "ollama":
                    summary = self._generate_with_ollama(prompt_template.format(text=documents[0].page_content),
                                                         progress_callback)
                else:
                    messages = [
                        SystemMessage(content="你是专业的文档分析师。You are a professional document analyst."),
                        HumanMessage(content=prompt_template.format(text=documents[0].page_content))
                    ]

                    summary = ""
                    chunk_count = 0

                    try:
                        if self.model_config["supports_streaming"]:
                            for chunk in self.llm.stream(messages):
                                if self.cancel_processing:
                                    return "用户已取消处理。Processing cancelled by user."

                                summary += chunk.content
                                chunk_count += 1

                                if progress_callback and chunk_count % 10 == 0:
                                    progress_callback(0.5 + 0.4 * min(chunk_count / 100, 1),
                                                      "生成摘要中... Generating summary...")
                        else:
                            response = self.llm.invoke(messages)
                            summary = response.content
                    except Exception as e:
                        return f"API调用失败 API call failed: {str(e)}"

            return self._format_summary(summary)

        except Exception as e:
            return f"摘要生成时出错 Error during summarization: {str(e)}"

    def _generate_with_ollama(self, prompt, progress_callback=None):
        """Generate text using Ollama"""
        try:
            if progress_callback:
                progress_callback(0.5, "使用Ollama生成摘要... Generating summary with Ollama...")

            # Create chat messages for Ollama
            messages = [
                SystemMessage(content="你是专业的文档分析师。You are a professional document analyst."),
                HumanMessage(content=prompt)
            ]

            summary = self.ollama_client.chat(
                messages,
                max_tokens=self.model_config["max_tokens"],
                temperature=self.model_config["temperature"]
            )

            if progress_callback:
                progress_callback(0.9, "Ollama摘要生成完成 Ollama summary generation complete")

            return summary
        except Exception as e:
            return f"Ollama生成失败 Ollama generation failed: {str(e)}"

    def _format_summary(self, summary: str) -> str:
        """Format summary for readability"""
        summary = re.sub(r'\n{3,}', '\n\n', summary)
        summary = re.sub(r'"\s*([^"]+)\s*"', r'"\1"', summary)
        return summary.strip()

    def analyze_document_structure(self, text: str) -> Dict[str, any]:
        """Quick document analysis"""
        analysis = {
            "total_words": len(text.split()),
            "total_characters": len(text),
            "total_sentences": text.count('.') + text.count('。'),
            "detected_language": "中文 Chinese" if len(
                re.findall(r'[\u4e00-\u9fff]', text[:1000])) > 100 else "英文 English",
            "text_quality": "损坏 Corrupted" if self.is_text_corrupted(text) else "良好 Good",
            "recommended_summary": "详细 detailed" if len(text.split()) > 5000 else "简洁 concise",
            "estimated_time": f"{max(1, len(text) // 10000)} 分钟 minutes",
            "current_model": self.model_config["name"]
        }
        return analysis

    def cancel_current_processing(self):
        """Cancel current processing operation"""
        self.cancel_processing = True


# CSS for the cat animation
CAT_ANIMATION_CSS = """
<style>
.cat-animation-container {
    height: 60px;
    position: relative;
    overflow: hidden;
    margin: 30px 0 20px 0;
}

.running-cat {
    position: absolute;
    width: 80px;
    height: 60px;
    animation: runCat 3s linear infinite;
}

@keyframes runCat {
    0% {
        left: -80px;
        transform: scaleX(1);
    }
    45% {
        left: 80%;
        transform: scaleX(1);
    }
    50% {
        left: 80%;
        transform: scaleX(-1);
    }
    95% {
        left: -80px;
        transform: scaleX(-1);
    }
    100% {
        left: -80px;
        transform: scaleX(1);
    }
}

.cat-body {
    width: 50px;
    height: 30px;
    background: #4F86F7;
    border-radius: 25px 25px 20px 20px;
    position: absolute;
    top: 20px;
    left: 15px;
}

.cat-head {
    width: 35px;
    height: 35px;
    background: #4F86F7;
    border-radius: 50%;
    position: absolute;
    top: 10px;
    left: 40px;
}

.cat-ear-left, .cat-ear-right {
    width: 0;
    height: 0;
    border-left: 8px solid transparent;
    border-right: 8px solid transparent;
    border-bottom: 15px solid #4F86F7;
    position: absolute;
    top: 5px;
}

.cat-ear-left {
    left: 38px;
    transform: rotate(-30deg);
}

.cat-ear-right {
    left: 55px;
    transform: rotate(30deg);
}

.cat-tail {
    width: 30px;
    height: 20px;
    background: #4F86F7;
    border-radius: 20px 0 0 20px;
    position: absolute;
    top: 15px;
    left: 0;
    transform-origin: right center;
    animation: wagTail 0.5s ease-in-out infinite;
}

@keyframes wagTail {
    0% {
        transform: rotate(-10deg);
    }
    50% {
        transform: rotate(10deg);
    }
    100% {
        transform: rotate(-10deg);
    }
}

.cat-leg {
    width: 8px;
    height: 15px;
    background: #4F86F7;
    position: absolute;
    bottom: 0;
    border-radius: 0 0 5px 5px;
    animation: runLegs 0.3s ease-in-out infinite;
}

.cat-leg1 { left: 20px; animation-delay: 0s; }
.cat-leg2 { left: 30px; animation-delay: 0.15s; }
.cat-leg3 { left: 45px; animation-delay: 0.1s; }
.cat-leg4 { left: 55px; animation-delay: 0.25s; }

@keyframes runLegs {
    0%, 100% {
        height: 15px;
    }
    50% {
        height: 10px;
    }
}

.cat-eye {
    width: 4px;
    height: 4px;
    background: white;
    border-radius: 50%;
    position: absolute;
    top: 18px;
}

.cat-eye-left { left: 50px; }
.cat-eye-right { left: 60px; }

.cat-whisker {
    width: 15px;
    height: 1px;
    background: #3366CC;
    position: absolute;
    top: 22px;
}

.cat-whisker1 { left: 65px; transform: rotate(10deg); }
.cat-whisker2 { left: 65px; transform: rotate(-10deg); }

/* Custom HTML output styling */
#summary_output {
    min-height: 400px;
    background-color: #f8f9fa;
    border: 1px solid #dee2e6;
    border-radius: 8px;
    overflow-y: auto;
    max-height: 600px;
}

/* Model selection styling */
.model-selector {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    border-radius: 10px;
    padding: 15px;
    margin: 10px 0;
    color: white;
}

.model-info {
    background: #f8f9fa;
    border-radius: 8px;
    padding: 10px;
    margin: 10px 0;
    border-left: 4px solid #007bff;
}
</style>
"""

# HTML for the cat animation
CAT_ANIMATION_HTML = """
<div class="cat-animation-container">
    <div class="running-cat">
        <div class="cat-tail"></div>
        <div class="cat-body"></div>
        <div class="cat-head"></div>
        <div class="cat-ear-left"></div>
        <div class="cat-ear-right"></div>
        <div class="cat-eye cat-eye-left"></div>
        <div class="cat-eye cat-eye-right"></div>
        <div class="cat-whisker cat-whisker1"></div>
        <div class="cat-whisker cat-whisker2"></div>
        <div class="cat-leg cat-leg1"></div>
        <div class="cat-leg cat-leg2"></div>
        <div class="cat-leg cat-leg3"></div>
        <div class="cat-leg cat-leg4"></div>
    </div>
</div>
"""


def create_optimized_gradio_interface():
    """Create the optimized Gradio interface with model selection"""

    summarizer = None
    is_processing = False

    def update_model_info(model_id):
        """Update model information display"""
        config = ModelConfig.get_config(model_id)
        if config:
            info_html = f"""
            <div class="model-info">
                <h4>🤖 模型信息 Model Information</h4>
                <p><strong>模型名称 Model:</strong> {config['name']}</p>
                <p><strong>提供商 Provider:</strong> {config['provider'].upper()}</p>
                <p><strong>API地址 API Base:</strong> <code>{config['api_base']}</code></p>
                <p><strong>流式输出 Streaming:</strong> {'✅ 支持 Supported' if config['supports_streaming'] else '❌ 不支持 Not Supported'}</p>
                <p><strong>获取API密钥 Get API Key:</strong> {config['api_key_help']}</p>
            </div>
            """
            return info_html, config['api_key_placeholder']
        return "", "Enter API key..."

    def set_api_key(model_id, api_key):
        """Initialize summarizer with selected model and API key"""
        nonlocal summarizer

        config = ModelConfig.get_config(model_id)
        if not config:
            return "❌ 无效的模型选择 Invalid model selection"

        # For Ollama, we don't need an API key
        if config["provider"] == "ollama":
            api_key = "dummy-key"  # Ollama doesn't need real API key
        elif not api_key or not api_key.strip():
            return "❌ 请输入有效的API密钥 Please enter a valid API key"

        try:
            summarizer = OptimizedDocumentSummarizer(model_id, api_key.strip())

            # Test connection
            connection_ok, message = summarizer.test_connection()
            if not connection_ok:
                return f"❌ 连接测试失败 Connection test failed: {message}"

            # OCR status
            ocr_status = "✅ OCR可用 OCR Available" if summarizer.ocr_available else "⚠️ OCR不可用 OCR Not Available"
            chinese_status = "✅ 中文OCR就绪 Chinese OCR Ready" if summarizer.chinese_ocr_available else "⚠️ 中文OCR未就绪 Chinese OCR Not Ready"

            return f"✅ {config['name']} 连接成功！Connected successfully! | {ocr_status} | {chinese_status}\n\n测试响应 Test Response: {message}"
        except Exception as e:
            return f"❌ 错误 Error: {str(e)}"

    def analyze_document(file):
        """Quick document analysis"""
        nonlocal summarizer

        if summarizer is None:
            return "❌ 请先设置模型和API密钥！Please set up model and API key first!"

        if file is None:
            return "❌ 请上传文件！Please upload a file!"

        try:
            # Quick text extraction for analysis
            text = summarizer.get_file_text(file.name, quality='fast', max_ocr_pages=1)

            if text.startswith("Error") or text.startswith("❌"):
                return text

            analysis = summarizer.analyze_document_structure(text)

            # Get file size
            file_size_mb = os.path.getsize(file.name) / (1024 * 1024)

            return f"""📊 **文档分析 Document Analysis:**

• **当前模型 Current Model:** {analysis['current_model']}
• **文件大小 File Size:** {file_size_mb:.2f} MB
• **总词数 Total Words:** {analysis['total_words']:,}
• **总字符 Total Characters:** {analysis['total_characters']:,}
• **检测语言 Detected Language:** {analysis['detected_language']}
• **文本质量 Text Quality:** {analysis['text_quality']}
• **推荐摘要类型 Recommended Summary:** {analysis['recommended_summary']}
• **预计处理时间 Estimated Time:** {analysis['estimated_time']}

⚠️ **重要限制 Important Limits:**
• 最大文本长度 Max text length: 200,000 字符 characters
• 最大文档块 Max chunks: 20
• API超时 API timeout: 60 秒 seconds
• 总处理超时 Total timeout: 5 分钟 minutes

💡 **性能提示 Performance Tips:**
• 大文档将自动截断 Large documents will be automatically truncated
• 使用"简洁"模式更快 Use 'Concise' mode for faster results
• 禁用OCR如果不需要 Disable OCR if not needed
• 考虑分割大文档 Consider splitting large documents
"""
        except Exception as e:
            return f"❌ 错误 Error: {str(e)}"

    def preview_text(file, use_ocr, ocr_language, quality, max_ocr_pages):
        """Preview extracted text"""
        nonlocal summarizer

        if summarizer is None:
            return "❌ 请先设置模型和API密钥！Please set up model and API key first!"

        if file is None:
            return "❌ 请上传文件！Please upload a file!"

        try:
            # Temporarily disable OCR if requested
            original_ocr_state = summarizer.ocr_available
            if not use_ocr:
                summarizer.ocr_available = False

            text = summarizer.get_file_text(
                file.name,
                ocr_language=ocr_language,
                quality=quality,
                max_ocr_pages=max_ocr_pages
            )

            # Restore OCR state
            summarizer.ocr_available = original_ocr_state

            if text.startswith("Error") or text.startswith("❌"):
                return text

            # Show preview (first 2000 characters)
            preview = text[:2000] + "..." if len(text) > 2000 else text

            return f"""📄 **文本预览 Text Preview:**

总长度 Total Length: {len(text)} 字符 characters
预计块数 Estimated Chunks: {len(summarizer.text_splitter.split_text(text))}
当前模型 Current Model: {summarizer.model_config['name']}

--- 预览 Preview ---
{preview}
"""
        except Exception as e:
            return f"❌ 错误 Error: {str(e)}"

    def process_document(file, summary_type, include_quotes, use_ocr, ocr_language,
                         output_language, quality, max_ocr_pages, progress=gr.Progress()):
        """Process document with progress tracking"""
        nonlocal summarizer, is_processing

        if summarizer is None:
            return "<div style='padding: 20px; color: red;'>❌ 请先设置模型和API密钥！Please set up model and API key first!</div>"

        if file is None:
            return "<div style='padding: 20px; color: red;'>❌ 请上传文件！Please upload a file!</div>"

        is_processing = True
        start_time = time.time()

        try:
            # Initial HTML with cat animation
            initial_html = f"""
            <div style='padding: 20px;'>
                {CAT_ANIMATION_HTML}
                <div style='text-align: center; margin-top: 20px;'>
                    <p>🤖 使用 {summarizer.model_config['name']} 开始处理...</p>
                    <p>🤖 Processing with {summarizer.model_config['name']}...</p>
                </div>
            </div>
            """

            # Progress callback that updates HTML
            current_progress_html = [initial_html]

            def update_progress(value, desc):
                elapsed = time.time() - start_time
                progress_html = f"""
                <div style='padding: 20px;'>
                    {CAT_ANIMATION_HTML}
                    <div style='text-align: center; margin-top: 20px;'>
                        <p>🤖 模型 Model: {summarizer.model_config['name']}</p>
                        <p>{desc} (已用时 Elapsed: {elapsed:.1f}s)</p>
                        <div style='width: 100%; background-color: #e0e0e0; border-radius: 5px; overflow: hidden; margin-top: 10px;'>
                            <div style='width: {value * 100}%; background-color: #4F86F7; height: 20px; transition: width 0.3s;'></div>
                        </div>
                        <p style='margin-top: 5px;'>{value * 100:.1f}%</p>
                    </div>
                </div>
                """
                current_progress_html[0] = progress_html
                progress(value, desc=desc)

            # Extract text
            update_progress(0.1, "开始提取文本... Starting text extraction...")

            # Temporarily disable OCR if requested
            original_ocr_state = summarizer.ocr_available
            if not use_ocr:
                summarizer.ocr_available = False

            text = summarizer.get_file_text(
                file.name,
                ocr_language=ocr_language,
                quality=quality,
                progress_callback=update_progress,
                max_ocr_pages=max_ocr_pages
            )

            # Restore OCR state
            summarizer.ocr_available = original_ocr_state

            if text.startswith("Error") or text.startswith("❌") or text.startswith("用户已取消"):
                is_processing = False
                return f"<div style='padding: 20px; color: red;'>{text}</div>"

            if len(text.strip()) < 10:
                is_processing = False
                return "<div style='padding: 20px; color: red;'>❌ 文档中未找到可读文本。No readable text found in the document.</div>"

            # Show text statistics
            update_progress(0.5, f"文本提取完成，长度: {len(text)} 字符 Text extracted, length: {len(text)} characters")

            # Generate summary
            update_progress(0.5,
                            f"使用 {summarizer.model_config['name']} 生成摘要... Generating summary with {summarizer.model_config['name']}...")

            # Create a progress tracking for streaming
            last_update_time = [time.time()]
            accumulated_text = [""]

            def streaming_update_progress(value, desc):
                current_time = time.time()
                # Update UI less frequently to avoid overwhelming
                if current_time - last_update_time[0] > 0.5:  # Update every 0.5 seconds
                    update_progress(value, desc)
                    last_update_time[0] = current_time

            summary = summarizer.summarize_text_streaming(
                text,
                summary_type,
                include_quotes,
                output_language,
                progress_callback=streaming_update_progress
            )

            elapsed_time = time.time() - start_time
            update_progress(1.0, f"完成！总用时: {elapsed_time:.1f}秒 Complete! Total time: {elapsed_time:.1f}s")

            # Add processing stats
            stats = f"\n\n---\n⏱️ 处理统计 Processing Stats:\n"
            stats += f"• 使用模型 Model used: {summarizer.model_config['name']}\n"
            stats += f"• 总用时 Total time: {elapsed_time:.1f} 秒 seconds\n"
            stats += f"• 文本长度 Text length: {len(text):,} 字符 characters\n"
            stats += f"• 文档块数 Document chunks: {len(summarizer.text_splitter.split_text(text))}\n"

            # Final result without cat animation
            final_html = f"""
            <div style='padding: 20px; white-space: pre-wrap; font-family: monospace;'>{summary + stats}</div>
            """

            is_processing = False
            return final_html

        except Exception as e:
            elapsed_time = time.time() - start_time
            is_processing = False
            return f"<div style='padding: 20px; color: red;'>❌ 错误 Error: {str(e)}\n⏱️ 失败时间 Failed after: {elapsed_time:.1f}秒 seconds</div>"

    def clear_cache():
        """Clear the document cache"""
        try:
            if summarizer and summarizer.cache:
                # Clear cache files
                cache_files = list(summarizer.cache.cache_path.glob("*.txt"))
                for f in cache_files:
                    try:
                        f.unlink()
                    except:
                        pass

                # Clear index
                summarizer.cache.index = {}
                summarizer.cache.save_index()

                return "✅ 缓存清除成功！Cache cleared successfully!"
        except Exception as e:
            return f"❌ 清除缓存时出错 Error clearing cache: {str(e)}"

    def cancel_processing():
        """Cancel current processing"""
        nonlocal summarizer
        if summarizer:
            summarizer.cancel_current_processing()
            return "<div style='padding: 20px; color: orange;'>⚠️ 已请求取消处理... Processing cancellation requested...</div>"
        return "<div style='padding: 20px;'>没有活动的处理可取消 No active processing to cancel</div>"

    def test_ollama_connection():
        """Test Ollama connection and list available models"""
        try:
            response = requests.get("http://localhost:11434/api/tags", timeout=5)
            if response.status_code == 200:
                models = response.json().get("models", [])
                model_names = [model["name"] for model in models]
                return f"✅ Ollama连接成功！Available models: {', '.join(model_names)}"
            else:
                return "❌ Ollama连接失败 Connection failed"
        except Exception as e:
            return f"❌ Ollama未运行或连接失败 Ollama not running or connection failed: {str(e)}"

    # Create the interface with custom CSS
    with gr.Blocks(title="多模型文档摘要生成器 Multi-Model Document Summarizer",
                   theme=gr.themes.Soft(),
                   css=CAT_ANIMATION_CSS) as interface:
        gr.Markdown(
            """
            # 🚀 多模型文档摘要生成器 Multi-Model Document Summarizer
            ## 支持 DeepSeek、豆包、Ollama 等多种AI模型 | Support for DeepSeek, Doubao, Ollama and more

            **🎯 支持的模型 Supported Models:**
            - 🚀 DeepSeek Chat - 深度求索的强大对话模型
            - 🌊 豆包 Pro/Lite - 字节跳动的智能对话助手
            - 🦙 Ollama - 本地运行的开源模型 (Llama3, Qwen, Mistral等)

            **✨ 新功能 New Features:**
            - 🔄 灵活的模型切换 Flexible model switching
            - 🏠 本地模型支持 Local model support
            - 🎨 更好的UI体验 Better UI experience
            - 📊 详细的处理统计 Detailed processing stats
            """
        )

        with gr.Row():
            with gr.Column(scale=1):
                gr.Markdown("### 🤖 模型配置 Model Configuration")

                # Model selection
                model_selector = gr.Dropdown(
                    choices=ModelConfig.get_model_choices(),
                    value="deepseek-chat",
                    label="选择AI模型 Select AI Model",
                    elem_classes=["model-selector"]
                )

                # Model info display
                model_info_display = gr.HTML(
                    value="",
                    label="模型信息 Model Information"
                )

                # API key input
                api_key_input = gr.Textbox(
                    label="API密钥 API Key",
                    placeholder="输入您的API密钥... Enter your API key...",
                    type="password"
                )

                with gr.Row():
                    api_key_button = gr.Button("🔑 设置模型 Setup Model", variant="primary")
                    test_ollama_button = gr.Button("🦙 测试Ollama Test Ollama", variant="secondary")

                api_key_status = gr.Textbox(label="连接状态 Connection Status", interactive=False, lines=4)

                # Cache control
                gr.Markdown("### 💾 缓存控制 Cache Control")
                clear_cache_button = gr.Button("🗑️ 清除缓存 Clear Cache", variant="secondary")
                cache_status = gr.Textbox(label="缓存状态 Cache Status", interactive=False)

            with gr.Column(scale=2):
                gr.Markdown("### 📤 文档上传 Document Upload")
                file_input = gr.File(
                    label="上传文档 Upload Document",
                    file_types=[".pdf", ".docx", ".doc", ".txt", ".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".gif"],
                    type="filepath"
                )

                with gr.Row():
                    analyze_button = gr.Button("📊 快速分析 Quick Analysis", variant="secondary")
                    preview_button = gr.Button("👁️ 预览文本 Preview Text", variant="secondary")

                analysis_output = gr.Markdown()

        with gr.Row():
            with gr.Column():
                summary_type = gr.Radio(
                    choices=[
                        ("📝 简洁 Concise (推荐 Recommended)", "concise"),
                        ("📖 详细 Detailed", "detailed"),
                        ("• 要点 Bullet Points", "bullet_points"),
                        ("💡 关键见解 Key Insights", "key_insights"),
                        ("📑 章节式 Chapter-wise", "chapter_wise")
                    ],
                    value="concise",
                    label="摘要类型 Summary Type"
                )

                # Performance settings
                gr.Markdown("### ⚡ 性能设置 Performance Settings")

                quality = gr.Radio(
                    choices=[
                        ("🚀 快速 Fast (100 DPI)", "fast"),
                        ("⚖️ 平衡 Balanced (150 DPI)", "balanced"),
                        ("🎯 高质量 High Quality (200 DPI)", "high")
                    ],
                    value="fast",
                    label="处理质量 Processing Quality"
                )

                max_ocr_pages = gr.Slider(
                    minimum=1,
                    maximum=50,
                    value=10,
                    step=1,
                    label="最大OCR页数 Maximum OCR Pages",
                    info="仅在启用OCR时使用 Only used when OCR is enabled"
                )

                include_quotes = gr.Checkbox(
                    label="包含引用 Include quotes",
                    value=False
                )

                use_ocr = gr.Checkbox(
                    label="🔍 启用OCR Enable OCR (扫描文档 for scanned docs)",
                    value=False
                )

                ocr_language = gr.Radio(
                    choices=[
                        ("自动检测 Auto-detect", "auto"),
                        ("中文优先 Chinese Priority", "chinese"),
                        ("仅英文 English Only", "english")
                    ],
                    value="auto",
                    label="OCR语言 OCR Language",
                    visible=False
                )

                output_language = gr.Radio(
                    choices=[
                        ("自动 Auto", "auto"),
                        ("中文 Chinese", "chinese"),
                        ("英文 English", "english")
                    ],
                    value="auto",
                    label="输出语言 Output Language"
                )

                with gr.Row():
                    summarize_button = gr.Button("🚀 生成摘要 Generate Summary", variant="primary", size="lg")
                    cancel_button = gr.Button("⏹️ 取消 Cancel", variant="stop", size="sm")

        gr.Markdown("### 📋 摘要输出 Summary Output")

        output_text = gr.HTML(
            label="摘要 Summary",
            value="",
            elem_id="summary_output"
        )

        # Event handlers
        model_selector.change(
            fn=update_model_info,
            inputs=[model_selector],
            outputs=[model_info_display, api_key_input]
        )

        api_key_button.click(
            fn=set_api_key,
            inputs=[model_selector, api_key_input],
            outputs=[api_key_status]
        )

        test_ollama_button.click(
            fn=test_ollama_connection,
            inputs=[],
            outputs=[api_key_status]
        )

        analyze_button.click(
            fn=analyze_document,
            inputs=[file_input],
            outputs=[analysis_output]
        )

        preview_button.click(
            fn=preview_text,
            inputs=[file_input, use_ocr, ocr_language, quality, max_ocr_pages],
            outputs=[analysis_output]
        )

        # Modified summarize button click handler
        def handle_summarize_click():
            # Show initial cat animation
            initial_html = f"""
            <div style='padding: 20px;'>
                <div style='text-align: center; margin-bottom: 40px;'>
                    <p>准备处理... Preparing to process...</p>
                </div>
                {CAT_ANIMATION_HTML}
            </div>
            """
            return initial_html

        summarize_button.click(
            fn=handle_summarize_click,
            inputs=[],
            outputs=[output_text]
        ).then(
            fn=process_document,
            inputs=[file_input, summary_type, include_quotes, use_ocr, ocr_language,
                    output_language, quality, max_ocr_pages],
            outputs=[output_text]
        )

        clear_cache_button.click(
            fn=clear_cache,
            inputs=[],
            outputs=[cache_status]
        )

        cancel_button.click(
            fn=cancel_processing,
            inputs=[],
            outputs=[output_text]
        )

        # Show/hide OCR language when OCR is toggled
        use_ocr.change(
            fn=lambda x: gr.update(visible=x),
            inputs=[use_ocr],
            outputs=[ocr_language]
        )

        # Initialize model info on load
        interface.load(
            fn=update_model_info,
            inputs=[model_selector],
            outputs=[model_info_display, api_key_input]
        )

        gr.Markdown(
            """
            ### 🚀 快速开始 Quick Start:

            1. **选择AI模型** Choose your AI model (DeepSeek/Doubao/Ollama)
            2. **设置API密钥** Set your API key (not needed for Ollama)
            3. **上传文档** Upload your document
            4. **点击"快速分析"** Click "Quick Analysis" to check document info
            5. **生成摘要** Click "Generate Summary"

            ### 🤖 模型说明 Model Instructions:

            **DeepSeek Chat:**
            - 获取API密钥：https://platform.deepseek.com/api_keys
            - 支持流式输出，速度较快
            - 适合大多数文档总结任务

            **豆包 (Doubao):**
            - 获取API密钥：https://console.volcengine.com/ark
            - 字节跳动开发，中文理解能力强
            - Pro版本功能更强，Lite版本速度更快

            **Ollama:**
            - 本地运行，无需API密钥
            - 先安装Ollama：https://ollama.ai
            - 启动服务：`ollama serve`
            - 下载模型：`ollama pull llama3` 或 `ollama pull qwen`

            ### ⚠️ 故障排除 Troubleshooting:

            - **连接失败**: 检查API密钥是否正确
            - **Ollama连接失败**: 确保Ollama服务正在运行
            - **处理超时**: 尝试使用更快的模型或减少文档大小
            - **OCR失败**: 检查是否安装了OCR依赖
            """
        )

    return interface


if __name__ == "__main__":
    print("""
    ====================================
    多模型文档摘要生成器
    MULTI-MODEL DOCUMENT SUMMARIZER
    ====================================

    支持的模型 Supported Models:
    - 🚀 DeepSeek Chat
    - 🌊 豆包 Doubao (ByteDance)
    - 🦙 Ollama (Local models)

    新功能 New Features:
    - 多模型选择 Multi-model selection
    - 本地模型支持 Local model support
    - 更好的UI体验 Better UI experience
    - 详细的处理统计 Detailed processing stats

    ====================================
    """)

    interface = create_optimized_gradio_interface()
    interface.launch(
        share=True,
        server_name="localhost",
        server_port=7860,
        show_error=True
    )